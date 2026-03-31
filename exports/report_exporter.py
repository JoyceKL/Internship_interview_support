from __future__ import annotations

from io import BytesIO
from typing import Dict, Any

from docx import Document


def build_markdown_report(cv_review: Dict[str, Any], qa: Dict[str, Any], audience: str = "lecturer") -> str:
    lines = [
        f"# CV Review Report ({audience})",
        f"- Student: {cv_review.get('student_info', {}).get('name', 'unknown')}",
        f"- Target role: {cv_review.get('target_role', '')}",
        f"- Overall score: {cv_review.get('cv_scores', {}).get('overall_score_100', 0)}/100",
        "",
        "## Strengths",
    ]
    for s in cv_review.get("strengths", []):
        lines.append(f"- {s}")

    lines.append("\n## Critical issues")
    for issue in cv_review.get("critical_issues", []):
        lines.append(f"- {issue}")

    lines.append("\n## Quick feedback")
    lines.append(cv_review.get("quick_feedback_for_student", ""))

    lines.append("\n## Interview Q&A")
    for idx, q in enumerate(qa.get("question_set", [])[:10], 1):
        lines.append(f"{idx}. **{q.get('question', '')}**")
        lines.append(f"   - Why: {q.get('why_this_is_asked', '')}")

    return "\n".join(lines)


def build_docx_report(cv_review: Dict[str, Any], qa: Dict[str, Any], audience: str = "lecturer") -> bytes:
    doc = Document()
    doc.add_heading(f"CV Review Report ({audience})", level=1)
    info = cv_review.get("student_info", {})
    doc.add_paragraph(f"Student: {info.get('name', 'unknown')}")
    doc.add_paragraph(f"Target role: {cv_review.get('target_role', '')}")
    doc.add_paragraph(f"Overall score: {cv_review.get('cv_scores', {}).get('overall_score_100', 0)}/100")

    doc.add_heading("Strengths", level=2)
    for s in cv_review.get("strengths", []):
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Critical issues", level=2)
    for s in cv_review.get("critical_issues", []):
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Interview Q&A", level=2)
    for q in qa.get("question_set", [])[:10]:
        doc.add_paragraph(q.get("question", ""), style="List Number")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
